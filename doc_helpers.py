from docx import Document

def expand_paragraphs(doc):
    """جلب كل الفقرات بما في ذلك الموجودة داخل الجداول."""
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def to_english_digits(text):
    """تحويل الأرقام العربية والفارسية (٠١٢٣... ۰۱۲۳...) إلى أرقام إنجليزية (0123)."""
    if not isinstance(text, str):
        text = str(text)
    # تشمل الأرقام العربية الشرقية (Hindi) والأرقام الفارسية
    arabic_indic_digits = "٠١٢٣٤٥٦٧٨٩۰۱۲۳۴۵۶۷۸۹"
    english_digits = "01234567890123456789"
    translation_table = str.maketrans(arabic_indic_digits, english_digits)
    return text.translate(translation_table)

def safe_replace_in_paragraph(paragraph, placeholders):
    """استبدال النصوص مع الحفاظ التام على التنسيق وضمان ظهور الأرقام بالإنجليزية."""
    for key, val in placeholders.items():
        if key not in paragraph.text:
            continue
            
        # تحويل القيمة لأرقام إنجليزية
        val_str = to_english_digits(str(val))
        
        # 1. الاستبدال في الـ runs المنفردة (هذا يحافظ على التنسيق تماماً)
        replaced_in_run = False
        for run in paragraph.runs:
            if key in run.text:
                new_text = run.text.replace(key, val_str)
                # خدعة المسافة لضمان الأرقام بالإنجليزية في وورد
                if any(c.isdigit() for c in val_str) and not new_text.startswith(" "):
                    new_text = " " + new_text
                run.text = new_text
                replaced_in_run = True
        
        # 2. الاستبدال الجراحي في حالة الكلمات المقطعة (Split Placeholders)
        # نقوم بدمج الـ runs المتأثرة فقط بدلاً من الفقرة كاملة لتجنب مشاكل التنسيق (مثل البولد)
        if not replaced_in_run:
            all_runs_text = [r.text for r in paragraph.runs]
            full_text = "".join(all_runs_text)
            
            if key in full_text:
                start_pos = full_text.find(key)
                end_pos = start_pos + len(key)
                
                current_len = 0
                start_run_idx = -1
                end_run_idx = -1
                
                for i, r_text in enumerate(all_runs_text):
                    run_start = current_len
                    run_end = current_len + len(r_text)
                    
                    if start_run_idx == -1 and run_start <= start_pos < run_end:
                        start_run_idx = i
                    if run_start < end_pos <= run_end:
                        end_run_idx = i
                        break
                    current_len = run_end
                
                if start_run_idx != -1 and end_run_idx != -1:
                    # دمج النص في نطاق الـ runs المتأثرة فقط
                    combined_text = "".join(all_runs_text[start_run_idx:end_run_idx+1])
                    new_combined = combined_text.replace(key, val_str)
                    
                    # إضافة مسافة بادئة للأرقام لضمان عرضها بالإنجليزية
                    if any(c.isdigit() for c in val_str) and not new_combined.startswith(" "):
                        new_combined = " " + new_combined
                    
                    # نحدث الـ run الأول في النطاق ونفرغ الباقي
                    paragraph.runs[start_run_idx].text = new_combined
                    for k in range(start_run_idx + 1, end_run_idx + 1):
                        paragraph.runs[k].text = ""

def safe_replace_in_doc(doc, placeholders):
    """تطبيق الاستبدال على كامل المستند."""
    for p in expand_paragraphs(doc):
        safe_replace_in_paragraph(p, placeholders)
