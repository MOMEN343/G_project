import random
import datetime
import os
import shutil
from docx import Document
from datetime import datetime
from datetime import date
from db import DataBase
from PyQt5 import uic,QtWidgets
from PyQt5.QtWidgets import QWidget
from PyQt5.QtWidgets import QMainWindow, QWidget, QPushButton, QLabel, QLineEdit, QTableWidget, QTableWidgetItem, QVBoxLayout, QMessageBox
from PyQt5.QtGui import QFontDatabase, QColor
from PyQt5.QtWidgets import QHBoxLayout
from PyQt5.QtWidgets import QHeaderView
from PyQt5.QtCore import  Qt, QPoint, QTimer
from PyQt5.QtWidgets import (
    QWidget, QVBoxLayout, QLabel, QPushButton, QScrollArea, QGridLayout,
    QTableWidget, QTableWidgetItem, QHeaderView, QDialog, QMessageBox,
    QMenu, QWidgetAction, QFrame, QHBoxLayout, QCheckBox
)

#-------------
class AdminWindow(QMainWindow):
    def __init__(self, main_shell=None):
        super().__init__()
        self.main_shell = main_shell
        self.db = DataBase()

        uic.loadUi("admin_dashboard.ui", self)

        # ربط الأزرار من الواجهة
        self.addEmployeeBtn.clicked.connect(self.open_add_user_window)
        self.logoutBtn.clicked.connect(self.log_out)
        
        self.addEmployeeBtn.setFocusPolicy(Qt.NoFocus)
        self.logoutBtn.setFocusPolicy(Qt.NoFocus)

        self.employeesTable.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        QFontDatabase.addApplicationFont("fonts/Alyamama-Bold.ttf")       
        self.setStyleSheet("""
    * {
        font-family: "Alyamama";
        color: white;
    }
""")

        self.db.cur.execute("""
            SELECT 
                username,
                password,
                full_name,
                email,
                phone,
                status,
                role.role_name
            FROM cms.users
            JOIN cms.role ON cms.users.role_id = cms.role.role_id
            """)

        result = self.db.cur.fetchall()
        for user in result: 
            row_position = self.employeesTable.rowCount()
            self.employeesTable.insertRow(row_position)

            self.employeesTable.setItem(row_position, 0, QTableWidgetItem((user[0])))
            self.employeesTable.setItem(row_position, 1, QTableWidgetItem((user[1])))
            self.employeesTable.setItem(row_position, 2, QTableWidgetItem((user[2])))
            self.employeesTable.setItem(row_position, 3, QTableWidgetItem((user[3])))
            self.employeesTable.setItem(row_position, 4, QTableWidgetItem((user[4])))
            self.employeesTable.setItem(row_position, 5, QTableWidgetItem((user[5])))
            self.employeesTable.setItem(row_position, 6, QTableWidgetItem(str(user[6])))


    def open_add_user_window(self):
        self.adduser_window = AddUserWindow(self)  #  مرّر الكائن نفسه
        self.adduser_window.show()


    def add_row(self, username, password,full_name,email,phone,status,role_id):
        row_position = self.employeesTable.rowCount()
        self.employeesTable.insertRow(row_position)

        self.employeesTable.setItem(row_position, 0, QTableWidgetItem(username))
        self.employeesTable.setItem(row_position, 1, QTableWidgetItem(password))
        self.employeesTable.setItem(row_position, 2, QTableWidgetItem(full_name))
        self.employeesTable.setItem(row_position, 3, QTableWidgetItem(email))
        self.employeesTable.setItem(row_position, 4, QTableWidgetItem(phone))
        self.employeesTable.setItem(row_position, 5, QTableWidgetItem(status))
        self.employeesTable.setItem(row_position, 6, QTableWidgetItem(role_id))

    def log_out (self):
        if self.main_shell:
            self.main_shell.switch_to_login()
        else:
            self.close()


class UserWindow(QMainWindow):
    def __init__(self, current_user_id, main_shell=None):
        super().__init__()
        self.current_user_id = current_user_id
        self.main_shell = main_shell
        self.db = DataBase()

        # Load the new UI
        uic.loadUi("employee.ui", self)
        
        # Apply the same logic as test.py for font application if needed, 
        # or rely on the UI file's stylesheet + global font loading in app.py
        self.setStyleSheet("""
        * {
            font-family: "Alyamama", "Segoe UI Symbol";
            color: #452829;
        }
        QLineEdit {
            border: 1px solid #ccc;
            border-radius: 5px;
            padding: 5px;
            background-color: white;
        }
        QLineEdit:focus {
            border: 1px solid #452829;
            background-color: #fcfcfc;
        }
        """)

        # Connect Buttons
        # Note: In the new UI, the buttons are named 'add_case' and 'docments' (typo in UI file acknowledged)
        self.add_case.clicked.connect(self.new_case_dialog)
        self.docments.clicked.connect(self.show_documents)
        self.logoutBtn.clicked.connect(self.log_out)
        self.master_record.clicked.connect(self.show_master_record)
        self.btn_scheduling.clicked.connect(self.show_scheduling)
        self.btn_save_session.clicked.connect(self.save_session)
        self.case2.clicked.connect(self.show_calendar)
        self.searchMasterRecord.textChanged.connect(self.filter_master_record)
        self.searchScheduling.textChanged.connect(self.filter_scheduling)
        
        if hasattr(self, 'notification'):
            self.notification.clicked.connect(self.show_notifications)
            self.notification.setFocusPolicy(Qt.NoFocus)
        
        # --- Notification Badge ---
        # The badge is now defined in the UI file as 'badge_label'
        # We need to re-parent it to the notification button to get the "overlay" effect
        if hasattr(self, 'notification') and hasattr(self, 'badge_label'):
            # Reparent to ensure it sits 'on top' or 'inside' the button's coordinate system
            self.badge_label.setParent(self.notification)
            self.badge_label.move(24, 2) # Adjust position to be top-right of bell
            
            # Timer to check for notifications
            self.timer = QTimer(self)
            self.timer.timeout.connect(self.update_badge)
            self.timer.start(5000) # Check every 5 seconds
            
            # Initial check
            self.update_badge()

        # Input validation styles


        # Ensure we start at the empty page
        if hasattr(self, 'mainStack'):
             self.mainStack.setCurrentWidget(self.page_empty)
             
        # Force alignment on the grid layout to prevent items from expanding to fill the whole area
        if hasattr(self, 'files_grid'):
            self.files_grid.setAlignment(Qt.AlignLeft | Qt.AlignTop)

        if hasattr(self, 'files_grid'):
            self.files_grid.setAlignment(Qt.AlignLeft | Qt.AlignTop)

    def update_badge(self):
        try:
            db = DataBase()
            db.cur.execute("""
                SELECT COUNT(*) FROM cms.notification 
                WHERE user_id = %s AND is_read = FALSE
            """, (self.current_user_id,))
            count = db.cur.fetchone()[0]
            db.close()

            if count > 0:
                self.badge_label.setText(str(count) if count < 10 else "9+")
                self.badge_label.show()
            else:
                self.badge_label.hide()
        except Exception as e:
            print(f"Error checking notifications: {e}")

    def show_notifications(self):
        # Fetch notifications - ONLY UNREAD
        db = DataBase()
        db.cur.execute("""
            SELECT notification_id, message, created_at, document_id 
            FROM cms.notification 
            WHERE user_id = %s AND is_read = FALSE
            ORDER BY created_at DESC
            LIMIT 10
        """, (self.current_user_id,))
        notifications = db.cur.fetchall()
        
        # Mark as read immediately when list is opened
        if notifications:
            ids = tuple([n[0] for n in notifications])
            if len(ids) == 1:
                db.cur.execute("UPDATE cms.notification SET is_read = TRUE WHERE notification_id = %s", (ids[0],))
            else:
                db.cur.execute("UPDATE cms.notification SET is_read = TRUE WHERE notification_id IN %s", (ids,))
            db.conn.commit()
        
        db.close()
        
        # Update badge immediately
        self.update_badge()

        menu = QMenu(self)
        menu.setMinimumWidth(350)
        menu.setStyleSheet("""
            QMenu {
                background-color: #452829;
                color: white;
                border: 1px solid #f3db93;
                border-radius: 10px;
                padding: 10px;
            }
        """)

        if not notifications:
            action = QWidgetAction(menu)
            lbl = QLabel("لا توجد إشعارات جديدة") # No new notifications
            lbl.setStyleSheet("color: #f3e8df; padding: 10px;")
            lbl.setAlignment(Qt.AlignCenter)
            action.setDefaultWidget(lbl)
            menu.addAction(action)
        else:
            for notif_id, msg, created_at, doc_id in notifications:
                # Format time HH:MM AM/PM
                time_str = created_at.strftime("%I:%M %p")
                
                # Custom Widget for Notification Item
                item_widget = QWidget()
                item_widget.setStyleSheet("background-color: transparent;")
                item_layout = QHBoxLayout(item_widget)
                item_layout.setContentsMargins(10, 5, 10, 5)
                item_layout.setDirection(QHBoxLayout.RightToLeft) # Name Right, Time Left
                
                # Message Label (Expanding)
                msg_label = QLabel(msg)
                msg_label.setStyleSheet("color: white; font-weight: bold; font-family: 'Alyamama'; font-size: 14px;")
                msg_label.setWordWrap(True)
                msg_label.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
                
                # Time Label (Fixed)
                time_label = QLabel(time_str)
                time_label.setStyleSheet("color: #f3db93; font-size: 12px;")
                time_label.setAlignment(Qt.AlignLeft | Qt.AlignVCenter)

                # Add widgets: msg stretches to fill space and push time to far left
                item_layout.addWidget(msg_label, 1)
                item_layout.addWidget(time_label, 0)
                
                action = QWidgetAction(menu)
                action.setDefaultWidget(item_widget)
                
                # Handle Click
                if doc_id:
                     # Connect triggered to a lambda that calls handle_click
                     # We use default arg d=doc_id to capture current value
                     action.triggered.connect(lambda checked, d=doc_id: self.handle_notification_click(d))
                
                menu.addAction(action)
                
                # Separator
                menu.addSeparator()

        # Show menu under the button
        menu.exec_(self.notification.mapToGlobal(QPoint(0, self.notification.height())))

    def handle_notification_click(self, document_id):
        self.show_documents(highlight_id=document_id)

    def reset_sidebar_styles(self):
        buttons = [self.add_case, self.docments, self.master_record, self.btn_scheduling, self.case2]
        for btn in buttons:
            btn.setFocusPolicy(Qt.NoFocus)
            btn.setProperty("active", False)
            btn.style().unpolish(btn)
            btn.style().polish(btn)

    def show_documents(self, highlight_id=None):
        self.reset_sidebar_styles()
        self.docments.setProperty("active", True)
        self.docments.style().unpolish(self.docments)
        self.docments.style().polish(self.docments)
        if hasattr(self, 'mainStack'):
            self.mainStack.setCurrentWidget(self.page_documents) # Show documents page

        # Mark all notifications for this user as read when entering documents page
        try:
            db_clear = DataBase()
            db_clear.cur.execute("UPDATE cms.notification SET is_read = TRUE WHERE user_id = %s", (self.current_user_id,))
            db_clear.conn.commit()
            db_clear.close()
            self.update_badge()
        except Exception as e:
            print(f"Error clearing notifications: {e}")

        # Clear existing items in the grid
        if hasattr(self, 'files_grid'):
            # First, try to fix the layout if it was set to Grid previously but we want a list behavior
            # straightforward way: use the grid as a vertical list (col 0 only)
            
            # Remove all widgets
            while self.files_grid.count():
                item = self.files_grid.takeAt(0)
                widget = item.widget()
                if widget:
                    widget.deleteLater()
            
            # Reset scaling or stretching if needed (optional but good practice)
            # self.files_grid.setColumnStretch(0, 1)

        db = DataBase()
        db.cur.execute("""
            SELECT d.file_path, d.document_id, n.created_at
            FROM cms.file_transfer ft
            JOIN cms.document d ON ft.document_id = d.document_id
            LEFT JOIN cms.notification n ON d.document_id = n.document_id AND n.user_id = ft.receiver_id
            WHERE ft.receiver_id = %s
            ORDER BY ft.transfer_date DESC
        """, (self.current_user_id,))
        files = db.cur.fetchall()
        db.close()

        row_idx = 0
        
        for (file_path, doc_id, created_at) in files:
            # Create a container widget for the row
            row_widget = QWidget()
            row_widget.setFixedHeight(80) # Fixed height for the row
            
            # Define Normal Style
            normal_style = """
                QWidget {
                    background-color: white;
                    border-bottom: 1px solid #e0e0e0;
                }
                QWidget:hover {
                    background-color: #f9f9f9;
                }
            """
            
            # Apply styling
            if highlight_id and doc_id == highlight_id:
                # Highlight Style
                highlight_style = normal_style + """
                QWidget {
                    background-color: #fff8e1;
                }
                """
                row_widget.setStyleSheet(highlight_style)
                
                # Auto-remove highlight after 3 seconds
                QTimer.singleShot(3000, lambda w=row_widget: w.setStyleSheet(normal_style))
                
            else:
                row_widget.setStyleSheet(normal_style)

            # Layout for the row
            layout = QHBoxLayout(row_widget)
            layout.setContentsMargins(20, 10, 20, 10)
            layout.setSpacing(15)

            # --- ELEMENTS ---
            
            # File Icon
            icon = QLabel("📄")
            icon.setStyleSheet("font-size: 30px; background: transparent; border: none;")
            
            # File Name
            file_name = os.path.basename(file_path)
            name_label = QLabel(file_name)
            name_label.setStyleSheet("font-size: 16px; font-weight: bold; color: #333; background: transparent; border: none;")
            name_label.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
            
            # Time Label (Replacing Document Type)
            # Format time as HH:MM AM/PM
            time_str = created_at.strftime("%I:%M %p") if created_at else ""
            time_label = QLabel(time_str)
            time_label.setStyleSheet("color: #777; font-size: 14px; background: transparent; border: none; font-family: 'Alyamama';")
            time_label.setAlignment(Qt.AlignRight | Qt.AlignVCenter)
            
            # Spacer to push buttons to the left
            spacer = QWidget()
            spacer.setSizePolicy(QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Preferred)
            spacer.setStyleSheet("background: transparent; border: none;")

            # Button: Extract Notification File
            btn_extract = QPushButton("استخراج ملف التبليغ")
            btn_extract.setCursor(Qt.PointingHandCursor)
            btn_extract.setFocusPolicy(Qt.NoFocus)
            btn_extract.setMinimumHeight(40)
            btn_extract.setStyleSheet("""
                QPushButton {
                    background-color: #452829; 
                    color: white; 
                    border-radius: 5px; 
                    padding: 5px 15px;
                    border: none;
                }
                QPushButton:hover {
                    background-color: #f3db93;
                    color: black;
                }
            """)
            btn_extract.clicked.connect(lambda checked, d=doc_id: self.extract_notification_file(d))

            # Button: Open
            btn_open = QPushButton("فتح")
            btn_open.setCursor(Qt.PointingHandCursor)
            btn_open.setFocusPolicy(Qt.NoFocus)
            btn_open.setMinimumHeight(40)
            btn_open.setStyleSheet("""
                QPushButton {
                    background-color: #452829; 
                    color: white; 
                    border-radius: 5px; 
                    padding: 5px 20px;
                    border: none;
                }
                QPushButton:hover {
                    background-color: #f3db93;
                    color: black;
                }
            """)
            btn_open.clicked.connect(lambda checked, p=file_path: self.open_file(p))

            # Add widgets to layout (Order for RTL: Right -> Left)
            layout.addWidget(icon)
            layout.addWidget(name_label)
            layout.addWidget(time_label)
            layout.addWidget(spacer) # Pushes subsequent items to the left
            layout.addWidget(btn_extract) # Will be to the left of spacer
            layout.addWidget(btn_open)    # Will be to the left of extract btn (furthest left)

            # Add row to grid (using it as a list)
            self.files_grid.addWidget(row_widget, row_idx, 0)
            row_idx += 1

        # Push everything to the top by adding a vertical spacer at the end
        if hasattr(self, 'files_grid'):
             # Create a spacer item that expands vertically
            vertical_spacer = QWidget()
            vertical_spacer.setSizePolicy(QtWidgets.QSizePolicy.Preferred, QtWidgets.QSizePolicy.Expanding)
            self.files_grid.addWidget(vertical_spacer, row_idx, 0)

    def extract_notification_file(self, doc_id):
        # Placeholder for extraction logic
        QMessageBox.information(self, "استخراج", f"جارٍ استخراج ملف التبليغ للمستند رقم {doc_id}...")

    def open_file(self, file_path):
        try:
            abs_path = os.path.abspath(file_path)
            if os.path.exists(abs_path):
                os.startfile(abs_path)
            else:
                QMessageBox.warning(self, "Error", f"File not found at:\n{abs_path}")
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Could not open file: {e}")
            
    def log_out(self):
        if self.main_shell:
            self.main_shell.switch_to_login()
        else:
            self.close()

    def show_calendar(self):
        # Highlight active button
        self.reset_sidebar_styles()
        self.case2.setProperty("active", True)
        self.case2.style().unpolish(self.case2)
        self.case2.style().polish(self.case2)
        
        self.mainStack.setCurrentWidget(self.page_calendar)
        
        # Set current date label
        self.label_calendar_date.setText(datetime.now().strftime("%B %d, %Y"))
        
        # 1. Fetch Judges from Database
        db = DataBase()
        db.cur.execute("SELECT full_name FROM cms.users WHERE role_id = 4") # role_id 4 = Judge
        judges_data = db.cur.fetchall()
        db.close()
        
        judge_names = [j[0] for j in judges_data]
        if not judge_names:
            judge_names = ["لا يوجد قضاة"]
        
        # Setup the calendar table (Schedule view)
        table = self.mainCalendarTable
        table.setColumnCount(len(judge_names))
        table.setHorizontalHeaderLabels(judge_names)
        table.verticalHeader().setVisible(True)
        table.setRowCount(11) # 09:00 to 19:00
        
        # ... existing setup ...
        hours = [f"{h:02d}:00" for h in range(9, 20)]
        table.setVerticalHeaderLabels(hours)
        table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        table.verticalHeader().setSectionResizeMode(QHeaderView.Stretch)

        # Clear existing items
        table.clearContents()

        # Add some dummy sessions for demonstration (mimicking the image)
        def add_event(row, col, text, color):
            item = QTableWidgetItem(text)
            item.setBackground(QColor(color))
            item.setForeground(QColor("white"))
            item.setTextAlignment(Qt.AlignCenter)
            font = item.font()
            font.setBold(True)
            item.setFont(font)
            table.setItem(row, col, item)

        # Demo events
        add_event(1, 0, "Check the plan\n10:00-11:00", "#4A90E2")
        add_event(1, 1, "Rawan, Busy\n10:00-11:00", "#9B9BEE")
        add_event(1, 2, "Ahmed, Busy\n10:00-11:00", "#A2D9CE")
        add_event(1, 3, "Yasser, Busy\n10:00-11:00", "#ABEBC6")
        add_event(1, 5, "Rawan, Busy\n09:00-11:00", "#EB984E")
        
        add_event(3, 1, "Rawan, Busy\n12:00-13:00", "#5DADE2")
        add_event(3, 2, "Ahmed, Busy\n12:00-13:00", "#16A085")
        add_event(3, 3, "Yasser, Busy\n12:00-13:00", "#2ECC71")
        
        add_event(4, 0, "Check the plan\n13:00-14:00", "#1A5276")
        add_event(4, 1, "Rawan, Busy", "#5499C7")

    def filter_master_record(self, text):
        """Filters the rows in the Master Record table based on search text."""
        table = self.masterRecordTable
        for row in range(table.rowCount()):
            match = False
            for col in range(table.columnCount()):
                item = table.item(row, col)
                if item and text.lower() in item.text().lower():
                    match = True
                    break
            table.setRowHidden(row, not match)

    def filter_scheduling(self, text):
        """Filters the rows in the Scheduling table based on search text."""
        table = self.schedulingTable
        for row in range(table.rowCount()):
            match = False
            for col in range(1, table.columnCount()): # Skip checkbox column
                item = table.item(row, col)
                if item and text.lower() in item.text().lower():
                    match = True
                    break
            table.setRowHidden(row, not match)

    def show_master_record(self):
        self.reset_sidebar_styles()
        self.master_record.setProperty("active", True)
        self.master_record.style().unpolish(self.master_record)
        self.master_record.style().polish(self.master_record)
        # 1. جلب بيانات سجل الأساس من قاعدة البيانات
        db = DataBase()
        db.cur.execute("""
            SELECT cc.case_id, c.plaintiff_name,c.defendant_name, ct.case_type, ct.filing_date, ct.status
            FROM cms.case_client cc
            JOIN cms.client c ON cc.client_id = c.client_id
            JOIN cms.court_case ct ON cc.case_id = ct.case_id
            ORDER BY ct.filing_date DESC
        """)
        records = db.cur.fetchall()
        db.close()

        # 2. تعبئة الجدول
        table = self.masterRecordTable
        table.verticalHeader().setVisible(False)
        table.setRowCount(0)
        table.setRowCount(len(records))
        
        # Adjust column sizes
        header = table.horizontalHeader()
        header.setSectionResizeMode(QHeaderView.Stretch)
        
        for row_idx, row_data in enumerate(records):
            for col_idx, value in enumerate(row_data):
                item = QTableWidgetItem(str(value))
                item.setTextAlignment(Qt.AlignCenter)
                table.setItem(row_idx, col_idx, item)
                
                # Special styling for status column
                if col_idx == 5: # Status column
                    if str(value) == "جديد":
                        item.setForeground(QColor("#2ECC71")) # Green
                    elif str(value) == "مغلق":
                        item.setForeground(QColor("#E74C3C")) # Red
                    item.setTextAlignment(Qt.AlignCenter)
                    font = item.font()
                    font.setBold(True)
                    item.setFont(font)
                
                table.setItem(row_idx, col_idx, item)

        # Clear search box when switching to this page
        self.searchMasterRecord.clear()
        
        # 3. عرض الصفحة في الـ stacked widget
        self.mainStack.setCurrentWidget(self.page_master_record)

    def show_scheduling(self):
        self.reset_sidebar_styles()
        self.btn_scheduling.setProperty("active", True)
        self.btn_scheduling.style().unpolish(self.btn_scheduling)
        self.btn_scheduling.style().polish(self.btn_scheduling)
        self.mainStack.setCurrentWidget(self.page_scheduling)
        # Populate table
        # Fetch cases that do NOT have a 'Scheduled' session
        db = DataBase()
        db.cur.execute("""
            SELECT cc.case_id, c.plaintiff_name, c.defendant_name, ct.case_type
            FROM cms.case_client cc
            JOIN cms.client c ON cc.client_id = c.client_id
            JOIN cms.court_case ct ON cc.case_id = ct.case_id
            WHERE cc.case_id NOT IN (
                SELECT case_id FROM cms.session WHERE status = 'Scheduled'
            )
            ORDER BY ct.filing_date DESC
        """)
        records = db.cur.fetchall()
        db.close()
        
        # Clear search box when switching to this page
        self.searchScheduling.clear()
        
        table = self.schedulingTable
        table.setRowCount(0)
        table.verticalHeader().setVisible(False)
        table.setRowCount(len(records))
        table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        
        self.scheduling_checkboxes = [] 
        
        for row, data in enumerate(records):
            # Checkbox in col 0
            chk = QCheckBox()
            cell_widget = QWidget()
            layout = QHBoxLayout(cell_widget)
            layout.addWidget(chk)
            layout.setAlignment(Qt.AlignCenter)
            layout.setContentsMargins(0,0,0,0)
            table.setCellWidget(row, 0, cell_widget)
            self.scheduling_checkboxes.append((chk, data[0])) # Store case_id
            
            # Data cols
            for col, val in enumerate(data, start=1):
                item = QTableWidgetItem(str(val))
                item.setTextAlignment(Qt.AlignCenter)
                table.setItem(row, col, item)

        # Populate Judge Combo
        self.judgeComboBox.clear()
        self.judgeComboBox.addItem("اختر القاضي")
        db = DataBase()
        db.cur.execute("SELECT user_id, full_name FROM cms.users WHERE role_id = 4")
        judges = db.cur.fetchall()
        db.close()
        for judge in judges:
            self.judgeComboBox.addItem(judge[1], judge[0])

    def save_session(self):
        selected_case_id = None
        if not hasattr(self, 'scheduling_checkboxes'):
             return

        for chk, case_id in self.scheduling_checkboxes:
            if chk.isChecked():
                selected_case_id = case_id
                break
        
        if not selected_case_id:
            QMessageBox.warning(self, "تنبيه", "الرجاء اختيار قضية")
            return
            
        # Get Judge ID
        judge_id = self.judgeComboBox.currentData()
        if not judge_id:
            QMessageBox.warning(self, "تنبيه", "الرجاء اختيار القاضي")
            return

        session_date = self.sessionDateInput.date().toString("yyyy-MM-dd")
        session_time = self.sessionTimeInput.time().toString("HH:mm")
        
        db = DataBase()
        try:
            # Check for conflict
            db.cur.execute("""
                SELECT count(*) FROM cms.session 
                WHERE judge_id = %s AND session_date = %s AND session_time = %s AND status = 'Scheduled'
            """, (judge_id, session_date, session_time))
            conflict_count = db.cur.fetchone()[0]
            
            if conflict_count > 0:
                QMessageBox.warning(self, "تنبيه", "يوجد جلسة أخرى لهذا القاضي في نفس الموعد!")
                return

            db.cur.execute("""
                INSERT INTO cms.session (session_date, session_time, status, case_id, judge_id)
                VALUES (%s, %s, %s, %s, %s)
            """, (session_date, session_time, 'Scheduled', selected_case_id, judge_id))
            db.conn.commit()

            # Get Case Number and actual file path
            db.cur.execute("SELECT case_number FROM cms.court_case WHERE case_id = %s", (selected_case_id,))
            case_number_val = db.cur.fetchone()[0]

            db.cur.execute("""
                SELECT file_path 
                FROM cms.document
                WHERE case_id = %s
                ORDER BY upload_date DESC LIMIT 1
            """, (selected_case_id,))
            res = db.cur.fetchone()
            
            if res:
                file_path = res[0]
                if os.path.exists(file_path):
                    doc = Document(file_path)
                    placeholders = {          
                        "{CASE_NUMBER}": str(case_number_val),
                        "{SESSION_DATE}": session_date,
                        "{SESSION_TIME}": session_time
                    }

                    def replace_placeholders(doc, placeholders):
                        for p in doc.paragraphs:
                            for run in p.runs:
                                for key, val in placeholders.items():
                                    if key in run.text:
                                        run.text = run.text.replace(key, val)
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for p in cell.paragraphs:
                                        for run in p.runs:
                                            for key, val in placeholders.items():
                                                if key in run.text:
                                                    run.text = run.text.replace(key, val)

                    replace_placeholders(doc, placeholders)
                    doc.save(file_path)

            QMessageBox.information(self, "نجاح", "تم حفظ الجلسة وتحديث ملف الدعوى بنجاح ✅")
            
            # Refresh the table
            self.show_scheduling()
            
        except Exception as e:
             QMessageBox.critical(self, "خطأ", f"حدث خطأ أثناء الحفظ: {e}")
        finally:
            db.close()


    # ================= إنشاء القضايا =================
    def new_case_dialog(self):
        db = DataBase()
        db.cur.execute("""
            SELECT client_id, plaintiff_name, defendant_name, case_type
            FROM cms.client
            WHERE client_id NOT IN (SELECT client_id FROM cms.case_client)
            """)
        clients = db.cur.fetchall()
        print(clients)
        db.close()

        if not clients:
            QMessageBox.information(self, "تنبيه", "لا توجد عملاء في قاعدة البيانات بعد.")
            return

        dialog = QDialog(self)
        dialog.setWindowTitle("إنشاء قضية جديدة")
        dialog.setFixedSize(550, 350)
        layout = QVBoxLayout(dialog)

        # واجهة من اليمين لليسار
        dialog.setLayoutDirection(Qt.RightToLeft)

        label = QLabel("اختر القضايا:")
        layout.addWidget(label)

        table = QTableWidget()
        table.setRowCount(len(clients))
        table.setColumnCount(4)
        table.setHorizontalHeaderLabels(["اختيار", "المدعي", "المدعى عليه", "نوع القضية"])
        table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        table.setLayoutDirection(Qt.RightToLeft)  # RTL للجدول

        checkboxes = []

        for row, client in enumerate(clients):
            client_id, plaintiff_name, defendant_name, case_type = client

            checkbox = QCheckBox()
            checkbox_widget = QWidget()
            cb_layout = QHBoxLayout(checkbox_widget)
            cb_layout.addWidget(checkbox)
            cb_layout.setAlignment(Qt.AlignCenter)
            cb_layout.setContentsMargins(0, 0, 0, 0)
            table.setCellWidget(row, 0, checkbox_widget)
            checkboxes.append(checkbox)

            table.setItem(row, 1, QTableWidgetItem(plaintiff_name))
            table.setItem(row, 2, QTableWidgetItem(defendant_name))
            table.setItem(row, 3, QTableWidgetItem(case_type))

        layout.addWidget(table)

        save_btn = QPushButton("إنشاء القضايا")
        layout.addWidget(save_btn)

        def create_case():
            selected_rows = [i for i, cb in enumerate(checkboxes) if cb.isChecked()]

            # i=> # الصف المحدد
            # cb=> # خانة الاختيار

            if not selected_rows:
                QMessageBox.warning(dialog, "تنبيه", "اختر قضية واحدة على الأقل")
                return

            db = DataBase()

            for idx in selected_rows:
                client_id, plaintiff_name, defendant_name, case_type = clients[idx]
                db.cur.execute("""
                    SELECT *
                    FROM cms.court_case
                """)
                case_count = len(db.cur.fetchall())
                case_number = f"{datetime.now().strftime('%Y')}/{case_count + 1}"
                filing_date = date.today()
                status = "مفتوحة"
                year = datetime.now().strftime('%Y')
                description = "-"

                db.cur.execute("""
                    INSERT INTO cms.court_case (case_type, case_number, status, filing_date, year, description, created_by)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                    RETURNING case_id
                """, (case_type, case_number, status, filing_date, year, description, self.current_user_id))
                new_case_id = db.cur.fetchone()[0]

                db.cur.execute("""
                    INSERT INTO cms.case_client (case_id, client_id, role_in_case)
                    VALUES (%s, %s, %s)
                """, (new_case_id, client_id, "Plaintiff"))

                db.cur.execute("""
                    UPDATE cms.document
                    SET case_id = %s
                    WHERE uploaded_by = %s AND case_id IS NULL AND document_type = %s
                """, (new_case_id, self.current_user_id, case_type))

            db.conn.commit()
            db.close()

            QMessageBox.information(dialog, "نجاح", f"تم إنشاء {len(selected_rows)} قضية بنجاح ✅")
            dialog.accept()


        save_btn.clicked.connect(create_case)
        dialog.exec_()


class AddUserWindow(QMainWindow):
    def __init__(self, admin_window):
        super().__init__()
        self.admin = admin_window
        uic.loadUi("add_user.ui", self)  # افترض وجود الواجهة

        # جلب البيانات من جدول role
        self.db = DataBase()
        self.role_combo.clear()
        self.db.cur.execute("SELECT role_id, role_name FROM cms.role")
        roles = self.db.cur.fetchall()
        for role in roles:
            role_id, role_name = role
            self.role_combo.addItem(role_name, role_id)
        self.setStyleSheet("""
        * {
            font-family: "Alyamama";
            color: #452829;
        }
    """)
        # ملء حالة الموظف في combo
        self.status_combo.clear()
        self.status_combo.addItem("اختر الحالة")  # placeholder
        self.status_combo.addItem("ACTIVE")
        self.status_combo.addItem("INACTIVE")

        # تكبير الحقول
        for widget in [self.username_input, self.password_input, self.full_name_input,
                       self.email_input, self.phone_input, self.status_combo, self.role_combo]:
            widget.setMinimumHeight(35)

        # ترتيب Status + Role جنب بعض
        horizontal_layout = QHBoxLayout()
        horizontal_layout.addWidget(self.status_combo)
        horizontal_layout.addWidget(self.role_combo)

        # نحذفهم من الـ layout القديم لو موجودين
        self.mainLayout.removeWidget(self.status_combo)
        self.mainLayout.removeWidget(self.role_combo)

        # نضيفهم قبل زر الحفظ
        index_save_btn = self.mainLayout.indexOf(self.saveBtn)
        self.mainLayout.insertLayout(index_save_btn, horizontal_layout)

        # ربط زر الحفظ
        self.saveBtn.clicked.connect(self.add_user_to_db)

    def add_user_to_db(self):
        username = self.username_input.text()
        password = self.password_input.text()
        full_name = self.full_name_input.text()
        email = self.email_input.text()
        phone = self.phone_input.text()

        # التحقق من اختيار الحالة
        if self.status_combo.currentIndex() == 0:
            QMessageBox.warning(self, "تحذير", "اختر حالة صالحة!")
            return
        status = self.status_combo.currentText()

        # جلب role_id من combo
        role_id = self.role_combo.currentData()

        # إدخال الموظف في قاعدة البيانات
        user_id = random.randint(20260000, 20269999)
        self.db.cur.execute(
            "INSERT INTO cms.users (user_id, username, password, full_name, email, phone, status, role_id) "
            "VALUES (%s,%s,%s,%s,%s,%s,%s,%s)",
            (user_id, username, password, full_name, email, phone, status, role_id)
        )
        self.db.conn.commit()
        self.db.close()

        # تحديث جدول الموظفين في الواجهة الرئيسية
        self.admin.add_row(username, password, full_name, email, phone, status, self.role_combo.currentText())

        QMessageBox.information(self, "نجاح", "تم إضافة المستخدم بنجاح")
        self.close()

class Petition_Clerks(QMainWindow):
    def __init__(self, current_user_id, main_shell=None):
        super().__init__()
        self.db = DataBase()
        self.c_u_i = current_user_id
        self.main_shell = main_shell
        self.current_case_data = None # To store selected case config

        # تحميل الواجهة من ملف UI
        uic.loadUi("petition_clerks2.ui", self)
        self.setStyleSheet("""
        * {
            font-family: "Alyamama";
            color: #452829;
        }
    """)
        
        # Connect main buttons
        self.sendFile.clicked.connect(self.process_full_workflow)
        self.logoutBtn.clicked.connect(self.log_out)
        self.sendFile.setFocusPolicy(Qt.NoFocus)
        self.logoutBtn.setFocusPolicy(Qt.NoFocus)
        
        # Case Configurations
        self.case_config = {
            "case1": {"label": "نفقة زوجة", "template": "لائحة دعوى نفقة زوجة.docx"},
            "case2": {"label": "عفش بيت", "template": "لائحة دعوى عفش بيت.docx"}, 
            "case3": {"label": "مهر مؤجل", "template": "لائحة دعوى مهر مؤجل.docx"},
            "case4": {"label": "نفقة عفش غيابي", "template": "nafqa.docx"},
            "case5": {"label": "نفقة زوجة غيابي", "template": "nafqa.docx"},
            "case6": {"label": "نفقة صغار", "template": "nafqa.docx"},
        }

        # Case Selection Buttons
        for btn_id in self.case_config.keys():
            if hasattr(self, btn_id):
                btn = getattr(self, btn_id)
                btn.clicked.connect(self.handle_case_selection_click)
                btn.setFocusPolicy(Qt.NoFocus)

        # Populate Receivers
        self.load_receivers()

        # Fix Tab Order
        self.setTabOrder(self.plaintiff_name, self.plaintiff_national_id)
        self.setTabOrder(self.plaintiff_national_id, self.plaintiff_phone)
        self.setTabOrder(self.plaintiff_phone, self.plaintiff_address)
        self.setTabOrder(self.plaintiff_address, self.defendant_name)
        self.setTabOrder(self.defendant_name, self.defendant_national_id)
        self.setTabOrder(self.defendant_national_id, self.defendant_phone)
        self.setTabOrder(self.defendant_phone, self.defendant_address)
        self.setTabOrder(self.defendant_address, self.comboBox)
        self.setTabOrder(self.comboBox, self.sendFile)

    def load_receivers(self):
        try:
            db = DataBase()
            db.cur.execute("SELECT user_id, full_name FROM cms.users WHERE role_id = '2'")
            users = db.cur.fetchall()
            self.comboBox.clear()
            self.comboBox.addItem("اختر الموظف المستلم...", None)
            for user_id, name in users:
                self.comboBox.addItem(name, user_id)
            db.close()
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to load receivers: {str(e)}")

    def handle_case_selection_click(self):
        btn = self.sender()
        if btn:
            btn_id = btn.objectName()
            if btn_id in self.case_config:
                config = self.case_config[btn_id]
                self.current_case_type = config["label"]
                self.current_template = config["template"]
                self.label_2.setText(f"أدخل بيانات لائحة دعوى {self.current_case_type}")

    def process_full_workflow(self):
        # Validation
        if not self.current_case_type:
            QMessageBox.warning(self, "تنبيه", "يرجى اختيار نوع القضية أولاً من القائمة الجانبية.")
            return

        receiver_index = self.comboBox.currentIndex()
        if receiver_index <= 0:
            QMessageBox.warning(self, "تنبيه", "يرجى اختيار الموظف المستلم.")
            return
        receiver_id = self.comboBox.itemData(receiver_index)

        # 1. Register Client
        try:
            # Re-open DB connection for transaction
            db = DataBase()
            db.cur.execute("""
                INSERT INTO cms.client (plaintiff_name, plaintiff_national_id, plaintiff_phone, plaintiff_address,
                                        defendant_name, defendant_national_id, defendant_phone, defendant_address, case_type)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING client_id
            """, (
                self.plaintiff_name.text(),
                self.plaintiff_national_id.text(),
                self.plaintiff_phone.text(),
                self.plaintiff_address.text(),
                self.defendant_name.text(),
                self.defendant_national_id.text(),
                self.defendant_phone.text(),
                self.defendant_address.text(),
                self.current_case_type 
            ))
            client_id = db.cur.fetchone()[0]
            db.conn.commit()
            db.close()
        except Exception as e:
             QMessageBox.warning(self, "Error", f"Failed to register client: {str(e)}")
             return

        # 2. Generate Document
        try:
            template_path = os.path.join("files", self.current_template)
            if not os.path.exists(template_path):
                QMessageBox.warning(self, "Error", f"Template file not found: {template_path}")
                return

            # Prepare filename
            safe_name = "".join([c for c in self.plaintiff_name.text() if c.isalnum() or c in (' ', '_')]).rstrip()
            final_filename = f"{self.current_case_type} - {safe_name}.docx"
            final_file = os.path.join("files", final_filename)
            
            shutil.copy2(template_path, final_file)
            doc = Document(final_file)

            # Date Mapping
            days_map = {
                "Monday": "الإثنين", "Tuesday": "الثلاثاء", "Wednesday": "الأربعاء",
                "Thursday": "الخميس", "Friday": "الجمعة", "Saturday": "السبت", "Sunday": "الأحد"
            }
            day_in_arabic = days_map.get(date.today().strftime("%A"), date.today().strftime("%A"))

            # Split address logic
            p_addr = self.plaintiff_address.text()
            p_from = p_addr.split('-')[0].strip() if '-' in p_addr else p_addr
            p_res = p_addr.split('-')[1].strip() if '-' in p_addr else "-"

            d_addr = self.defendant_address.text()
            d_from = d_addr.split('-')[0].strip() if '-' in d_addr else d_addr
            d_res = d_addr.split('-')[1].strip() if '-' in d_addr else "-"

            placeholders = {
                "{DATE_DAY}": day_in_arabic,
                "{DATE_FULL}": date.today().strftime("%d/%m/%Y"),
                "{PLAINTIFF_NAME}": self.plaintiff_name.text(),
                "{PLAINTIFF_FROM}": p_from,
                "{PLAINTIFF_RESIDENT}": p_res,
                "{DEFENDANT_NAME}": self.defendant_name.text(),
                "{DEFENDANT_FROM}": d_from,
                "{DEFENDANT_RESIDENT}": d_res,            
            }

            def replace_in_doc(doc, placeholders, p_from, p_res, d_from, d_res):
                import re
                
                # 1. Multi-pass Placeholder Replacement (Robust against run splitting)
                paragraphs = list(doc.paragraphs)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            paragraphs.extend(list(cell.paragraphs))
                
                for p in paragraphs:
                    full_text = p.text
                    # A. Standard Placeholder {}
                    updated_tags = False
                    for key, val in placeholders.items():
                        if key in full_text:
                            full_text = full_text.replace(key, str(val))
                            updated_tags = True
                    
                    # B. Literal Anchors with Underscores (based on user's manual template)
                    updated_anchors = False
                    
                    # Normalized identification (Kashida-agnostic)
                    # Checking for Plaintiff and Defendant with kashida support
                    is_plaintiff_line = re.search(r"المدع[ـ]*ية", full_text)
                    is_defendant_line = re.search(r"المدع[ـ]*ي[ـ]*[\s]*عليه|المدع[ـ]*ى[ـ]*[\s]*عليه", full_text)
                    
                    # Handle Plaintiff Line (Heading or Signature)
                    if is_plaintiff_line and not is_defendant_line:
                        # Only replace if the name is not already there (prevents duplication with {})
                        if self.plaintiff_name.text() not in full_text:
                            full_text = re.sub(r"المدع[ـ]*ية\s*/?[ـ_\s]*", f"المدعية/ {self.plaintiff_name.text()} ", full_text)
                        
                        if p_from not in full_text:
                            full_text = re.sub(r"من[ـ_\s]+", f"من {p_from} ", full_text)
                        if p_res not in full_text:
                            full_text = re.sub(r"وسكان[ـ_\s]+", f"وسكان {p_res} ", full_text)
                        updated_anchors = True
                    
                    # Handle Defendant Line
                    if is_defendant_line:
                        # Only replace if the name is not already there
                        if self.defendant_name.text() not in full_text:
                            full_text = re.sub(r"(المدع[ـ]*ي[ـ]*\s*عليه|المدع[ـ]*ى[ـ]*\s*عليه)\s*/?[ـ_\s]*", f"المدعى عليه/ {self.defendant_name.text()} ", full_text)
                        
                        if d_from not in full_text:
                            full_text = re.sub(r"من[ـ_\s]+", f"من {d_from} ", full_text)
                        if d_res not in full_text:
                            full_text = re.sub(r"وسكان[ـ_\s]+", f"وسكان {d_res} ", full_text)
                        updated_anchors = True

                    if updated_tags or updated_anchors:
                        if len(p.runs) > 0:
                            p.runs[0].text = full_text
                            for i in range(1, len(p.runs)):
                                p.runs[i].text = ""
                        else:
                            p.add_run(full_text)

            # Execution
            replace_in_doc(doc, placeholders, p_from, p_res, d_from, d_res)
            doc.save(final_file)
            
             # Save to DB
            db = DataBase()
            case_id = None
            db.cur.execute("""
                INSERT INTO cms.document (document_type, file_path, uploaded_by,case_id)
                VALUES (%s, %s, %s, %s)
                RETURNING document_id
            """, (self.current_case_type, final_file, self.c_u_i, case_id))
            self.document_id = db.cur.fetchone()[0]
            db.conn.commit()
            db.close()
            
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to generate file: {str(e)}")
            return

        # 3. Send File (File Transfer)
        try:
            db = DataBase()
            transfer_id = random.randint(1, 1000000)
            transfer_date = date.today()
            status = "pending"

            db.cur.execute("""
                INSERT INTO cms.file_transfer (transfer_id, transfer_date, status, document_id, sender_id, receiver_id)
                VALUES (%s, %s, %s, %s, %s, %s)
            """, (transfer_id, transfer_date, status, self.document_id, self.c_u_i, receiver_id))

            db.conn.commit()
            db.close()
        except Exception as e:
            QMessageBox.warning(self, "Error", f"Failed to send file: {str(e)}")
            return

        # 4. Create Notification
        try:
            notification_msg = f"({self.current_case_type} - \"{self.plaintiff_name.text().strip()}\" جديد)"
            db = DataBase()
            db.cur.execute("""
                INSERT INTO cms.notification (message, user_id, created_at, document_id)
                VALUES (%s, %s, NOW(), %s)
            """, (notification_msg, receiver_id, self.document_id))
            db.conn.commit()
            db.close()
        except Exception as e:
            print(f"Failed to send notification: {e}") # Non-blocking error

        msg = f"تم تسجيل الموكل، إنشاء الملف وإرساله بنجاح!\n\nالمسار:\n{final_file}"
        QMessageBox.information(self, "نجاح", msg)
        
        # Clear fields
        self.plaintiff_name.clear()
        self.plaintiff_national_id.clear()
        self.plaintiff_phone.clear()
        self.plaintiff_address.clear()
        self.defendant_name.clear()
        self.defendant_national_id.clear()
        self.defendant_phone.clear()
        self.defendant_address.clear()
        self.comboBox.setCurrentIndex(0)
        self.label_2.setText("أدخل بيانات لائحة الدعوى:")

    def log_out (self):
        if self.main_shell:
            self.main_shell.switch_to_login()
        else:
            self.close()  