from docx import Document
import re
import os

def test_replacement(file_path):
    doc = Document(file_path)
    
    placeholders = {
        "{PLAINTIFF_NAME}": "اسم المدعية التجريبي",
        "{PLAINTIFF_ADDRESS}": "عنوان المدعية التجريبي",
        "{DEFENDANT_NAME}": "اسم المدعى عليه التجريبي",
        "{DEFENDANT_ADDRESS}": "عنوان المدعى عليه التجريبي",
        "{DATE_FULL}": "30/01/2026"
    }
    
    pl_name = placeholders.get("{PLAINTIFF_NAME}", "")
    pl_addr = placeholders.get("{PLAINTIFF_ADDRESS}", "")
    df_name = placeholders.get("{DEFENDANT_NAME}", "")
    df_addr = placeholders.get("{DEFENDANT_ADDRESS}", "")

    def process_paragraph(p):
        handled = False
        text = p.text
        if "المدعية/" in text or "المدعي عليه/" in text or "المدعي عليه " in text:
            if "ــــ" in text:
                if "المدعية/" in text:
                    repls = [pl_name, pl_addr, pl_addr]
                else:
                    repls = [df_name, df_addr, df_addr]
                
                count = 0
                def repl_func(m):
                    nonlocal count
                    res = repls[count] if count < len(repls) else m.group(0)
                    count += 1
                    return res
                
                new_text = re.sub(r"ـ+", repl_func, text)
                for key, val in placeholders.items():
                    new_text = new_text.replace(key, val)
                
                p.text = new_text
                handled = True

        if not handled:
            for run in p.runs:
                for key, val in placeholders.items():
                    if key in run.text:
                        run.text = run.text.replace(key, val)

    for p in doc.paragraphs:
        process_paragraph(p)
        
    # Show results
    print("--- Results ---")
    for i, p in enumerate(doc.paragraphs):
        if "التجريبي" in p.text:
            print(f"P{i}: {p.text}")

if __name__ == "__main__":
    import sys
    sys.stdout.reconfigure(encoding='utf-8')
    template = r"c:\Users\kraiz\G_project_clean\files\لائحة دعوى نفقة زوجة.docx"
    test_replacement(template)
