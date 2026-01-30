from docx import Document
import os

def inspect_docx(file_path):
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return

    doc = Document(file_path)
    print(f"--- Inspecting: {file_path} ---")
    
    print("\n[PARAGRAPHS]")
    for i, p in enumerate(doc.paragraphs):
        # We search for underscores or keywords
        if "المدعية" in p.text or "المدعى" in p.text or "ــــ" in p.text:
            print(f"P{i}: {repr(p.text)}")
            for j, run in enumerate(p.runs):
                print(f"  R{j}: {repr(run.text)}")

    print("\n[TABLES]")
    for i, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, p in enumerate(cell.paragraphs):
                    if "المدعية" in p.text or "المدعى" in p.text or "ــــ" in p.text:
                        print(f"T{i} R{r_idx} C{c_idx} P{p_idx}: {repr(p.text)}")
                        for run in p.runs:
                            print(f"  Run: {repr(run.text)}")

if __name__ == "__main__":
    import sys
    # Ensure stdout handles UTF-8
    sys.stdout.reconfigure(encoding='utf-8')
    template = r"c:\Users\kraiz\G_project_clean\files\نفقة زوجة - فلان الفلاني.docx"
    inspect_docx(template)
